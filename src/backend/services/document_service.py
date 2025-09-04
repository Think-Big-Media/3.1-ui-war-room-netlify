"""
Document processing service for War Room platform.
Handles file upload, text extraction, chunking, and vector embedding.
"""

import os
import hashlib
import asyncio
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple, BinaryIO
from pathlib import Path
import aiofiles
import aiofiles.os
from sqlalchemy.orm import Session
from sqlalchemy import and_

# Text extraction libraries
import PyPDF2
import pandas as pd
from docx import Document as DocxDocument
import magic

# Text processing
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter

from models.document import Document, DocumentChunk, DocumentType, ProcessingStatus
from core.database import get_db
from core.config import settings
from core.pinecone_config import pinecone_manager
from utils.text_processing import clean_text, extract_keywords

logger = logging.getLogger(__name__)


class DocumentService:
    """
    Service for handling document operations including upload, processing, and search.

    Features:
    - Multi-format file upload (PDF, CSV, DOCX, TXT)
    - Text extraction and cleaning
    - Intelligent text chunking
    - Vector embedding generation
    - Pinecone storage and retrieval
    - Organization-level data isolation
    """

    def __init__(self):
        self.upload_directory = Path("uploads")
        self.upload_directory.mkdir(exist_ok=True)

        # Text splitter for creating chunks
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.DOCUMENT_CHUNK_SIZE,
            chunk_overlap=settings.DOCUMENT_CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    async def upload_document(
        self,
        file: BinaryIO,
        filename: str,
        organization_id: str,
        user_id: str,
        title: Optional[str] = None,
        description: Optional[str] = None,
        tags: Optional[List[str]] = None,
        db: Session = None,
    ) -> Optional[Document]:
        """
        Upload and process a document.

        Args:
            file: File binary data
            filename: Original filename
            organization_id: Organization ID
            user_id: Uploading user ID
            title: Optional document title
            description: Optional description
            tags: Optional tags list
            db: Database session

        Returns:
            Document model or None if failed
        """
        try:
            # Read file content
            file_content = await file.read()
            file_size = len(file_content)

            # Validate file size
            max_size = settings.MAX_DOCUMENT_SIZE_MB * 1024 * 1024
            if file_size > max_size:
                logger.error(f"File too large: {file_size} bytes (max: {max_size})")
                return None

            # Generate file hash for deduplication
            file_hash = hashlib.sha256(file_content).hexdigest()

            # Detect MIME type
            mime_type = magic.from_buffer(file_content, mime=True)

            # Determine document type
            document_type = self._get_document_type(filename, mime_type)
            if not document_type:
                logger.error(f"Unsupported file type: {mime_type}")
                return None

            # Check for duplicate
            existing_doc = (
                db.query(Document)
                .filter(
                    and_(
                        Document.file_hash == file_hash,
                        Document.organization_id == organization_id,
                        Document.is_active == True,
                    )
                )
                .first()
            )

            if existing_doc:
                logger.info(f"Duplicate document detected: {filename}")
                return existing_doc

            # Generate unique filename
            file_ext = Path(filename).suffix.lower()
            unique_filename = f"{hashlib.md5(f'{filename}{datetime.utcnow()}'.encode()).hexdigest()}{file_ext}"
            file_path = self.upload_directory / organization_id / unique_filename

            # Create directory if needed
            file_path.parent.mkdir(parents=True, exist_ok=True)

            # Save file to disk
            async with aiofiles.open(file_path, "wb") as f:
                await f.write(file_content)

            # Create document record
            document = Document(
                organization_id=organization_id,
                uploaded_by=user_id,
                filename=unique_filename,
                original_filename=filename,
                file_path=str(file_path),
                file_size=file_size,
                file_hash=file_hash,
                mime_type=mime_type,
                document_type=document_type,
                title=title or filename,
                description=description,
                tags=tags or [],
                processing_status=ProcessingStatus.PENDING,
            )

            db.add(document)
            db.commit()
            db.refresh(document)

            # Start processing asynchronously
            asyncio.create_task(self._process_document_async(document.id))

            logger.info(f"Document uploaded successfully: {document.id}")
            return document

        except Exception as e:
            logger.error(f"Document upload failed: {str(e)}")
            if db:
                db.rollback()
            return None

    async def _process_document_async(self, document_id: str) -> None:
        """Process document asynchronously in background."""
        try:
            async for db in get_db():
                await self.process_document(document_id, db)
                break
        except Exception as e:
            logger.error(f"Async document processing failed: {str(e)}")

    async def process_document(self, document_id: str, db: Session) -> bool:
        """
        Process document: extract text, chunk, and create embeddings.

        Args:
            document_id: Document ID to process
            db: Database session

        Returns:
            Success status
        """
        try:
            # Get document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                logger.error(f"Document not found: {document_id}")
                return False

            # Update status to processing
            document.processing_status = ProcessingStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            db.commit()

            # Extract text based on document type
            extracted_text = await self._extract_text(document)
            if not extracted_text:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "Failed to extract text from document"
                db.commit()
                return False

            # Clean and process text
            cleaned_text = clean_text(extracted_text)

            # Create text chunks
            text_chunks = self.text_splitter.split_text(cleaned_text)

            if not text_chunks:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "No text chunks created from document"
                db.commit()
                return False

            # Store in vector database
            metadata = {
                "document_id": document.id,
                "filename": document.original_filename,
                "document_type": document.document_type,
                "title": document.title,
                "tags": document.tags,
                "mime_type": document.mime_type,
                "uploaded_at": document.created_at.isoformat(),
            }

            success = await pinecone_manager.upsert_document(
                document_id=document.id,
                text_chunks=text_chunks,
                metadata=metadata,
                organization_id=document.organization_id,
            )

            if not success:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = "Failed to store in vector database"
                db.commit()
                return False

            # Create chunk records in database
            for i, chunk_text in enumerate(text_chunks):
                vector_id = f"{document.id}_chunk_{i}"

                chunk = DocumentChunk(
                    document_id=document.id,
                    organization_id=document.organization_id,
                    chunk_index=i,
                    chunk_text=chunk_text,
                    chunk_length=len(chunk_text),
                    vector_id=vector_id,
                    keywords=extract_keywords(chunk_text),
                )
                db.add(chunk)

            # Update document with processing results
            document.processing_status = ProcessingStatus.COMPLETED
            document.processing_completed_at = datetime.utcnow()
            document.is_vectorized = True
            document.vector_count = len(text_chunks)
            document.extracted_text = cleaned_text
            document.text_length = len(cleaned_text)
            document.search_keywords = " ".join(extract_keywords(cleaned_text))

            db.commit()

            logger.info(
                f"Document processed successfully: {document_id} ({len(text_chunks)} chunks)"
            )
            return True

        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            if document:
                document.processing_status = ProcessingStatus.FAILED
                document.processing_error = str(e)
                db.commit()
            return False

    async def _extract_text(self, document: Document) -> Optional[str]:
        """
        Extract text from document based on type.

        Args:
            document: Document model

        Returns:
            Extracted text or None
        """
        try:
            file_path = Path(document.file_path)

            if document.document_type == DocumentType.PDF:
                return await self._extract_pdf_text(file_path)
            elif document.document_type == DocumentType.CSV:
                return await self._extract_csv_text(file_path)
            elif document.document_type == DocumentType.DOCX:
                return await self._extract_docx_text(file_path)
            elif document.document_type == DocumentType.TXT:
                return await self._extract_txt_text(file_path)
            else:
                logger.error(f"Unsupported document type: {document.document_type}")
                return None

        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return None

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""

        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

        return text.strip()

    async def _extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV file."""
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)

            # Convert to text representation
            text = f"CSV Data Summary:\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += f"Rows: {len(df)}\n\n"

            # Add column descriptions
            for col in df.columns:
                text += f"Column '{col}':\n"
                text += f"  Data type: {df[col].dtype}\n"
                text += f"  Non-null values: {df[col].count()}\n"

                if df[col].dtype in ["object"]:
                    unique_values = df[col].unique()[:10]  # First 10 unique values
                    text += f"  Sample values: {', '.join(map(str, unique_values))}\n"
                elif df[col].dtype in ["int64", "float64"]:
                    text += f"  Min: {df[col].min()}, Max: {df[col].max()}, Mean: {df[col].mean():.2f}\n"

                text += "\n"

            # Add sample data
            text += "Sample Data (first 5 rows):\n"
            text += df.head().to_string(index=False)

            return text

        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            return ""

    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(row_text) + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""

    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as file:
                return await file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    async with aiofiles.open(file_path, "r", encoding=encoding) as file:
                        return await file.read()
                except UnicodeDecodeError:
                    continue

            logger.error(f"Could not decode text file: {file_path}")
            return ""

    def _get_document_type(
        self, filename: str, mime_type: str
    ) -> Optional[DocumentType]:
        """Determine document type from filename and MIME type."""
        ext = Path(filename).suffix.lower()

        # PDF
        if ext == ".pdf" or "pdf" in mime_type:
            return DocumentType.PDF

        # CSV
        if ext == ".csv" or "csv" in mime_type:
            return DocumentType.CSV

        # DOCX
        if ext in [".docx", ".doc"] or "document" in mime_type:
            return DocumentType.DOCX

        # Text
        if ext == ".txt" or "text" in mime_type:
            return DocumentType.TXT

        return None

    async def search_documents(
        self,
        query: str,
        organization_id: str,
        user_id: str,
        filters: Optional[Dict[str, Any]] = None,
        top_k: int = 5,
        db: Session = None,
    ) -> List[Dict[str, Any]]:
        """
        Search documents using vector similarity.

        Args:
            query: Search query
            organization_id: Organization ID
            user_id: Searching user ID
            filters: Additional filters
            top_k: Number of results
            db: Database session

        Returns:
            List of search results
        """
        try:
            start_time = datetime.utcnow()

            # Search in Pinecone
            results = await pinecone_manager.search_documents(
                query=query,
                organization_id=organization_id,
                top_k=top_k,
                filter_metadata=filters,
            )

            search_duration = (datetime.utcnow() - start_time).total_seconds() * 1000

            # Log search for analytics
            # search_log = DocumentSearch(
            #     user_id=user_id,
            #     organization_id=organization_id,
            #     query=query,
            #     results_count=len(results),
            #     top_score=results[0]["score"] if results else None,
            #     avg_score=sum(r["score"] for r in results) / len(results) if results else None,
            #     search_duration_ms=int(search_duration),
            #     result_document_ids=[r["document_id"] for r in results],
            #     applied_filters=filters or {}
            # )
            # db.add(search_log)
            # db.commit()

            logger.info(
                f"Document search completed: {len(results)} results in {search_duration:.0f}ms"
            )
            return results

        except Exception as e:
            logger.error(f"Document search failed: {str(e)}")
            return []

    async def delete_document(
        self, document_id: str, organization_id: str, db: Session
    ) -> bool:
        """
        Delete document and all associated data.

        Args:
            document_id: Document ID
            organization_id: Organization ID (for security)
            db: Database session

        Returns:
            Success status
        """
        try:
            # Get document
            document = (
                db.query(Document)
                .filter(
                    and_(
                        Document.id == document_id,
                        Document.organization_id == organization_id,
                    )
                )
                .first()
            )

            if not document:
                logger.error(f"Document not found: {document_id}")
                return False

            # Delete from vector database
            if document.is_vectorized:
                await pinecone_manager.delete_document(document_id, organization_id)

            # Delete file from disk
            try:
                await aiofiles.os.remove(document.file_path)
            except FileNotFoundError:
                logger.warning(f"File not found on disk: {document.file_path}")

            # Soft delete document (keep for audit trail)
            document.is_active = False
            document.deleted_at = datetime.utcnow()

            db.commit()

            logger.info(f"Document deleted successfully: {document_id}")
            return True

        except Exception as e:
            logger.error(f"Document deletion failed: {str(e)}")
            return False


# Global service instance
document_service = DocumentService()


# Helper functions
async def upload_document_file(
    file: BinaryIO, filename: str, organization_id: str, user_id: str, **kwargs
) -> Optional[Document]:
    """Upload and process document file."""
    async for db in get_db():
        return await document_service.upload_document(
            file, filename, organization_id, user_id, db=db, **kwargs
        )


async def search_organization_documents(
    query: str, organization_id: str, user_id: str, **kwargs
) -> List[Dict[str, Any]]:
    """Search documents within organization."""
    async for db in get_db():
        return await document_service.search_documents(
            query, organization_id, user_id, db=db, **kwargs
        )


async def remove_document(document_id: str, organization_id: str) -> bool:
    """Remove document and all data."""
    async for db in get_db():
        return await document_service.delete_document(document_id, organization_id, db)
